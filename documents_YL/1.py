from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Заголовок документа', 0)

document.add_paragraph('Абзац без форматирования')

# тут у нас будет более сложный абзац
p = document.add_paragraph('Часть абзаца обычным текстом, ')
p.add_run('часть жирным шрифтом, ').bold = True
p.add_run(' а часть ')
p.add_run('курсивом.').italic = True

document.add_heading('Заголовок первого уровня', level=1)
document.add_paragraph('Некоторая цитата', style='Intense Quote')

document.add_paragraph('Элемент ненумерованного списка',
                       style='List Bullet')
document.add_paragraph('Элемент нумерованного списка',
                       style='List Number')

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Номер'
hdr_cells[1].text = 'Название'
hdr_cells[2].text = 'Количество'

document.save('test.docx')